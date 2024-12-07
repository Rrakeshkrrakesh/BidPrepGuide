import streamlit as st
import google.generativeai as genai
from google.generativeai.types import GenerationConfig
import json
import os
from PyPDF2 import PdfReader
from docx import Document
from pytesseract import image_to_string
from PIL import Image

# Set your Gemini API Key as an environment variable (Recommended for Streamlit Cloud)
# In Streamlit Cloud, set this in the "Secrets" section of your app settings.
os.environ["API_KEY"] = st.secrets["API_KEY"]  # Access the API key from secrets



# Initialize Gemini client 
genai.configure(api_key=os.environ['API_KEY'])
gemini_model = genai.GenerativeModel("gemini-1.5-flash") # Or your preferred model




# --- Function definitions ---
def process_uploaded_file(uploaded_file):
    try:
        file_content = uploaded_file.read()
        if uploaded_file.name.endswith(".pdf"):
            reader = PdfReader(file_content)
            extracted_text = " ".join(page.extract_text() for page in reader.pages)
        elif uploaded_file.name.endswith(".docx"):
            doc = Document(file_content)
            extracted_text = " ".join([p.text for p in doc.paragraphs])
        elif uploaded_file.name.endswith(".txt"):
            extracted_text = file_content.decode("utf-8")
        else:  # Assume it's an image
            image = Image.open(file_content)
            extracted_text = image_to_string(image)
        return extracted_text
    except Exception as e:
        st.error(f"An error occurred during file processing: {e}")  # Streamlit error message
        return None

def extract_entities_with_gemini(text):
    try:
        response = gemini_model.generate_content(  # Using the global gemini_model
            f"""Extract the following entities from this text, providing them in a structured JSON format:\n
            * Dates (submission deadlines, important dates)\n
            * Eligibility Criteria\n
            * Technical Specifications\n
            * Pricing and Financial information\n
            * Scope of Work\n
            * Penalty Clauses\n
            * Contact Information\n\n
            Text: {text}""",
          generation_config=GenerationConfig(
                temperature=0.0,  # Example: Setting temperature for more deterministic outputs
                candidate_count=1 
          )
        )

        try:
            extracted_entities = json.loads(response.text)
            return extracted_entities
        except json.JSONDecodeError:
            st.warning("Gemini's response could not be parsed as strict JSON.  Returning raw text.") # Streamlit warning
            return response.text

    except Exception as e:
        st.error(f"An error occurred during entity extraction: {e}") # Streamlit error message
        return None

def determine_eligibility(extracted_info):
  # ... (Your eligibility logic - you'll need to implement this)
  return "Unknown" #Placeholder

def chat_with_document(text, user_message):
    try:

        chat = gemini_model.start_chat(context=text) #Using global gemini_model
        response = chat.send_message(user_message)
        return response.text

    except Exception as e:
        st.error(f"An error occurred during chat: {e}") # Streamlit error message
        return None

def generate_bid_document(data, output_file="Generated_Bid_Document.docx"):
  try:
        doc = Document()
        doc.add_heading("Bid Submission", level=1)
        doc.add_paragraph(f"Bidder Name: {data.get('name', 'N/A')}")
        doc.add_paragraph(f"Eligibility Status: {data.get('eligibility', 'N/A')}")
        doc.add_paragraph(f"Estimated Cost: ${data.get('cost', 'N/A')}")
        # ... Add all other required bid data from the 'data' dictionary
        doc.save(output_file)


        with open(output_file, "rb") as file: #Provide download link for the generated file
            st.download_button(
                label="Download Bid Document",
                data=file,
                file_name=output_file,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document" #or mime type of your file
            )
   except Exception as e:
        st.error(f"An error occurred during bid document generation: {e}")


# --- Streamlit UI ---
st.title("Bid Preparation App")

uploaded_file = st.file_uploader("Upload Bid Document", type=["pdf", "docx", "txt", "jpg", "png"]) # Added image types


if uploaded_file is not None:
    extracted_text = process_uploaded_file(uploaded_file)
    if extracted_text:
        with st.spinner("Analyzing document..."):
            extracted_info = extract_entities_with_gemini(extracted_text)

            if extracted_info:
                st.header("Extracted Information")
                st.write(extracted_info)


                st.header("Bid Preparation")
                bidder_name = st.text_input("Enter your name/company:")
                estimated_cost = st.number_input("Enter estimated cost (in USD):", value=0.0, step=1.0)
                # ... get other necessary bid data from user using streamlit widgets

                if st.button("Generate Bid Document"):
                    bid_data = { #Add more keys as required
                        "name": bidder_name,
                        "eligibility": determine_eligibility(extracted_info),
                        "cost": estimated_cost
                    }
                    generate_bid_document(bid_data)
                    st.success("Bid document generated!")



                st.header("Chat with the Document")
                user_message = st.text_input("Enter your question:")

                if st.button("Ask"):

                    response = chat_with_document(extracted_text, user_message)
                    if response:
                        st.write("Gemini's Response:")
                        st.write(response)


    else:
        st.error("Could not extract text from the document.")
