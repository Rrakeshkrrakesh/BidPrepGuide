from PyPDF2 import PdfReader
from docx import Document
from pytesseract import image_to_string
from PIL import Image
import gemini

# Gemini API Initialization
gemini_api_key = "your_gemini_api_key"  # Replace with your actual API key
gemini_client = gemini.Client(api_key=AIzaSyDsO43WG0w0Y9bOCahrgwgMjMlkT9pYFWg)

def process_uploaded_file(file):
    """Process uploaded document and extract text."""
    if file.name.endswith(".pdf"):
        reader = PdfReader(file)
        return " ".join(page.extract_text() for page in reader.pages)
    elif file.name.endswith(".docx"):
        doc = Document(file)
        return " ".join([p.text for p in doc.paragraphs])
    elif file.name.endswith(".txt"):
        return file.read().decode("utf-8")
    else:
        return image_to_string(Image.open(file))

def extract_entities_with_gemini(text):
    """Extract entities using Gemini API."""
    response = gemini_client.named_entity_recognition(text)
    return response.get("entities", {})

def check_eligibility(user_profile, criteria):
    """Compare user profile with extracted criteria."""
    return {key: user_profile.get(key, None) == value for key, value in criteria.items()}

def generate_bid_document(data):
    """Generate a bid document with user data."""
    doc = Document()
    doc.add_heading("Bid Submission", level=1)
    doc.add_paragraph(f"Bidder Name: {data['name']}")
    doc.add_paragraph(f"Eligibility Status: {data['eligibility']}")
    doc.add_paragraph(f"Estimated Cost: ${data['cost']}")
    doc.save("Generated_Bid_Document.docx")
